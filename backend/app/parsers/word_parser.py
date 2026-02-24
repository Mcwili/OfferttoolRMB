"""
Word-Parser
Verwendet python-docx für Word-Dokumente
"""

from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches
from io import BytesIO
import re
from app.models.project import ProjectFile


class WordParser:
    """Parser für Word-Dokumente (.docx)"""
    
    async def parse(self, file_content: bytes, file_obj: ProjectFile) -> Dict[str, Any]:
        """
        Extrahiert Daten aus Word-Datei
        Returns: Dict mit extrahierten Entitäten
        """
        doc = Document(BytesIO(file_content))
        
        source_info = {
            "datei": file_obj.original_filename,
            "datei_id": file_obj.id,
            "upload_am": file_obj.upload_date.isoformat() if file_obj.upload_date else None
        }
        
        extracted_data = {
            "raeume": [],
            "anlagen": [],
            "geraete": [],
            "anforderungen": [],
            "termine": [],
            "leistungen": [],
            "raw_tables": [],
            "full_text": "",
            "metadata": {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "image_count": len(doc.inline_shapes),
                "sections": []
            }
        }
        
        # Strukturierte Abschnitte erkennen
        sections = self._extract_sections(doc)
        extracted_data["metadata"]["sections"] = [
            {"name": name, "paragraph_indices": indices} 
            for name, indices in sections.items()
        ]
        
        # Volltext extrahieren
        full_text_parts = []
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                full_text_parts.append(text)
        extracted_data["full_text"] = "\n".join(full_text_parts)
        
        # Anforderungen aus Textabschnitten erkennen
        anforderungen = self._extract_anforderungen(doc, source_info, sections)
        extracted_data["anforderungen"].extend(anforderungen)
        
        # Tabellen extrahieren
        tables_data = self._extract_tables(doc, source_info)
        extracted_data["raeume"].extend(tables_data.get("raeume", []))
        extracted_data["geraete"].extend(tables_data.get("geraete", []))
        extracted_data["anlagen"].extend(tables_data.get("anlagen", []))
        extracted_data["termine"].extend(tables_data.get("termine", []))
        extracted_data["leistungen"].extend(tables_data.get("leistungen", []))
        
        # IMMER: Alle Tabellen als Rohdaten extrahieren (auch wenn Typ nicht erkannt wurde)
        for table_idx, table in enumerate(doc.tables):
            raw_table = self._extract_raw_table(table, table_idx, source_info)
            if raw_table:
                extracted_data["raw_tables"].append(raw_table)
        
        # Listen extrahieren
        list_anforderungen = self._extract_lists(doc, source_info)
        extracted_data["anforderungen"].extend(list_anforderungen)
        
        return extracted_data
    
    def _extract_sections(self, doc: Document) -> Dict[str, List[int]]:
        """Extrahiert Abschnittsstruktur basierend auf Überschriften"""
        sections = {}
        current_section = None
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            style_name = paragraph.style.name.lower()
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # Überschriften erkennen
            if "heading" in style_name or "überschrift" in style_name:
                level = 1
                if "heading 1" in style_name or "überschrift 1" in style_name:
                    level = 1
                elif "heading 2" in style_name or "überschrift 2" in style_name:
                    level = 2
                elif "heading 3" in style_name or "überschrift 3" in style_name:
                    level = 3
                
                current_section = text.lower()
                if current_section not in sections:
                    sections[current_section] = []
                sections[current_section].append(para_idx)
        
        return sections
    
    def _extract_anforderungen(self, doc: Document, source_info: Dict[str, Any], sections: Dict[str, List[int]]) -> List[Dict[str, Any]]:
        """Extrahiert Anforderungen aus Word-Dokument"""
        anforderungen = []
        
        # Schlüsselwörter für Anforderungen
        requirement_keywords = [
            "luftwechsel", "temperatur", "feuchtigkeit", "anforderung", "vorgabe",
            "muss", "soll", "sollte", "erforderlich", "notwendig", "benötigt",
            "luftqualität", "raumklima", "komfort", "energieeffizienz"
        ]
        
        # Kategorien erkennen
        categories = {
            "technisch": ["luftwechsel", "temperatur", "feuchtigkeit", "luftqualität", "raumklima"],
            "organisatorisch": ["termin", "abgabe", "freigabe", "koordination"],
            "energie": ["energieeffizienz", "energie", "verbrauch", "leistung"]
        }
        
        current_section = None
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            text_lower = text.lower()
            
            # Abschnitt erkennen
            for section_name, para_indices in sections.items():
                if para_idx in para_indices:
                    current_section = section_name
                    break
            
            # Anforderung erkennen
            if any(kw in text_lower for kw in requirement_keywords):
                # Priorität bestimmen
                prioritaet = "mittel"
                if any(kw in text_lower for kw in ["muss", "erforderlich", "notwendig", "kritisch"]):
                    prioritaet = "hoch"
                elif any(kw in text_lower for kw in ["sollte", "empfohlen", "optional"]):
                    prioritaet = "niedrig"
                
                # Kategorie bestimmen
                kategorie = None
                for cat, keywords in categories.items():
                    if any(kw in text_lower for kw in keywords):
                        kategorie = cat
                        break
                if not kategorie:
                    kategorie = "allgemein"
                
                # SIA-Phase erkennen (falls erwähnt)
                sia_phase = None
                sia_patterns = [
                    r"sia\s*(\d+)",
                    r"phase\s*(\d+)",
                    r"projektierung",
                    r"vorprojekt",
                    r"bauprojekt"
                ]
                for pattern in sia_patterns:
                    match = re.search(pattern, text_lower)
                    if match:
                        if "projektierung" in text_lower:
                            sia_phase = "SIA 103 - Projektierung"
                        elif "vorprojekt" in text_lower:
                            sia_phase = "SIA 104 - Vorprojekt"
                        elif "bauprojekt" in text_lower:
                            sia_phase = "SIA 105 - Bauprojekt"
                        elif match.group(1):
                            sia_phase = f"SIA {match.group(1)}"
                        break
                
                anforderung = {
                    "id": f"REQ_{para_idx:04d}",
                    "beschreibung": text,
                    "kategorie": kategorie,
                    "prioritaet": prioritaet,
                    "sia_phase": sia_phase,
                    "quelle": {
                        **source_info,
                        "absatz": para_idx,
                        "abschnitt": current_section
                    }
                }
                anforderungen.append(anforderung)
        
        return anforderungen
    
    def _extract_tables(self, doc: Document, source_info: Dict[str, Any]) -> Dict[str, List[Dict[str, Any]]]:
        """Extrahiert Tabellen aus Word-Dokument"""
        tables_data = {
            "raeume": [],
            "geraete": [],
            "anlagen": [],
            "termine": [],
            "leistungen": []
        }
        
        for table_idx, table in enumerate(doc.tables):
            # Erste Zeile als Header verwenden
            if len(table.rows) < 2:
                continue
            
            header_row = table.rows[0]
            headers = [cell.text.strip().lower() for cell in header_row.cells]
            
            # Tabellentyp erkennen
            header_text = " ".join(headers)
            
            # Raumliste
            if any(kw in header_text for kw in ["raum", "room", "fläche", "flaeche"]):
                for row_idx, row in enumerate(table.rows[1:], start=1):
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) < len(headers):
                        continue
                    
                    raum_data = {}
                    for idx, header in enumerate(headers):
                        if idx < len(cells):
                            if any(kw in header for kw in ["raum", "room", "nummer"]):
                                raum_data["name"] = cells[idx]
                            elif any(kw in header for kw in ["fläche", "flaeche", "m²"]):
                                try:
                                    raum_data["flaeche_m2"] = float(cells[idx].replace(",", "."))
                                except ValueError:
                                    pass
                            elif any(kw in header for kw in ["nutzung", "art"]):
                                raum_data["nutzungsart"] = cells[idx]
                    
                    if raum_data.get("name"):
                        raum = {
                            "id": f"Raum_{raum_data['name'].replace(' ', '_')}_{table_idx}_{row_idx}",
                            "name": raum_data["name"],
                            "flaeche_m2": raum_data.get("flaeche_m2"),
                            "nutzungsart": raum_data.get("nutzungsart"),
                            "quelle": {
                                **source_info,
                                "tabelle": table_idx,
                                "zeile": row_idx
                            }
                        }
                        tables_data["raeume"].append(raum)
            
            # Geräteliste
            elif any(kw in header_text for kw in ["geraet", "equipment", "gerät", "typ"]):
                for row_idx, row in enumerate(table.rows[1:], start=1):
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) < len(headers):
                        continue
                    
                    geraet_data = {}
                    for idx, header in enumerate(headers):
                        if idx < len(cells):
                            if any(kw in header for kw in ["geraet", "equipment", "gerät", "name"]):
                                geraet_data["name"] = cells[idx]
                            elif any(kw in header for kw in ["typ", "type"]):
                                geraet_data["typ"] = cells[idx]
                            elif any(kw in header for kw in ["leistung", "power", "kw"]):
                                try:
                                    geraet_data["leistung_kw"] = float(cells[idx].replace(",", "."))
                                except ValueError:
                                    pass
                    
                    if geraet_data.get("name") or geraet_data.get("typ"):
                        geraet = {
                            "id": f"GER_{table_idx}_{row_idx}",
                            "typ": geraet_data.get("typ") or geraet_data.get("name"),
                            "name": geraet_data.get("name"),
                            "leistung_kw": geraet_data.get("leistung_kw"),
                            "quelle": {
                                **source_info,
                                "tabelle": table_idx,
                                "zeile": row_idx
                            }
                        }
                        tables_data["geraete"].append(geraet)
            
            # Terminplan
            elif any(kw in header_text for kw in ["datum", "termin", "date", "deadline"]):
                for row_idx, row in enumerate(table.rows[1:], start=1):
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) < len(headers):
                        continue
                    
                    termin_data = {}
                    for idx, header in enumerate(headers):
                        if idx < len(cells):
                            if any(kw in header for kw in ["beschreibung", "bezeichnung", "aufgabe"]):
                                termin_data["beschreibung"] = cells[idx]
                            elif any(kw in header for kw in ["datum", "termin", "date"]):
                                termin_data["termin_datum"] = cells[idx]
                    
                    if termin_data.get("beschreibung"):
                        termin = {
                            "id": f"TERM_{table_idx}_{row_idx}",
                            "beschreibung": termin_data["beschreibung"],
                            "termin_datum": termin_data.get("termin_datum"),
                            "quelle": {
                                **source_info,
                                "tabelle": table_idx,
                                "zeile": row_idx
                            }
                        }
                        tables_data["termine"].append(termin)
        
        return tables_data
    
    def _extract_lists(self, doc: Document, source_info: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extrahiert Listen und Aufzählungen"""
        anforderungen = []
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Listen erkennen (Bullet Points oder nummerierte Listen)
            if paragraph.style.name.startswith("List") or paragraph.style.name.startswith("Aufzählung"):
                # Prüfe ob es eine Anforderung ist
                requirement_keywords = [
                    "anforderung", "vorgabe", "muss", "soll", "erforderlich"
                ]
                
                if any(kw in text.lower() for kw in requirement_keywords):
                    anforderung = {
                        "id": f"REQ_LIST_{para_idx:04d}",
                        "beschreibung": text,
                        "kategorie": "allgemein",
                        "prioritaet": "mittel",
                        "quelle": {
                            **source_info,
                            "absatz": para_idx,
                            "typ": "liste"
                        }
                    }
                    anforderungen.append(anforderung)
        
        return anforderungen
    
    def _extract_raw_table(self, table, table_idx: int, source_info: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """Extrahiert eine Tabelle als Rohdaten (auch wenn Typ nicht erkannt wurde)"""
        if len(table.rows) < 1:
            return None
        
        # Header extrahieren
        header_row = table.rows[0]
        headers = [cell.text.strip() if cell.text.strip() else f"Spalte_{i+1}" 
                   for i, cell in enumerate(header_row.cells)]
        
        # Datenzeilen extrahieren
        rows_data = []
        for row_idx, row in enumerate(table.rows[1:], start=1):
            row_dict = {}
            for col_idx, cell in enumerate(row.cells):
                if col_idx < len(headers):
                    row_dict[headers[col_idx]] = cell.text.strip()
            
            # Überspringe leere Zeilen
            if any(value for value in row_dict.values()):
                rows_data.append(row_dict)
        
        if not rows_data:
            return None
        
        return {
            "table_index": table_idx,
            "headers": headers,
            "rows": rows_data,
            "row_count": len(rows_data),
            "column_count": len(headers),
            "quelle": source_info
        }