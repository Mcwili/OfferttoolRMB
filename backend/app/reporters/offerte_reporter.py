"""
Ingenieurleistungs-Offerte Reporter
Generiert RMB-Offerte basierend auf Projekt-Daten
"""

from typing import Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from pathlib import Path
from app.models.project import Project


class OfferteReporter:
    """Reporter für Ingenieurleistungs-Offerte"""
    
    def __init__(self):
        self.output_dir = Path("reports")
        self.output_dir.mkdir(exist_ok=True)
    
    async def generate(self, project: Project, project_data: Dict[str, Any]) -> str:
        """
        Generiert Offerte als Word-Dokument
        Returns: Pfad zur generierten Datei
        """
        doc = Document()
        
        # Titel
        title = doc.add_heading(f"Ingenieurleistungs-Offerte: {project.name}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Projektinformationen
        doc.add_heading("1. Projektinformationen", 1)
        projekt = project_data.get("projekt", {})
        p = doc.add_paragraph()
        p.add_run("Projektname: ").bold = True
        p.add_run(projekt.get("name", "-"))
        
        p = doc.add_paragraph()
        p.add_run("Standort: ").bold = True
        p.add_run(projekt.get("standort", "-"))
        
        p = doc.add_paragraph()
        p.add_run("Erstellt am: ").bold = True
        p.add_run(datetime.now().strftime("%d.%m.%Y"))
        
        # Leistungsbeschreibung nach SIA-Phasen
        doc.add_heading("2. Leistungsbeschreibung", 1)
        
        sia_phasen = [
            ("SIA 103 - Projektierung", "Projektanalyse, Raumanalyse, Konzeptentwicklung"),
            ("SIA 104 - Vorprojekt", "Vorprojekt, Kostenplanung, Terminplanung"),
            ("SIA 105 - Bauprojekt", "Bauprojekt, Ausschreibung, Vergabe")
        ]
        
        for phase, beschreibung in sia_phasen:
            doc.add_heading(f"2.{sia_phasen.index((phase, beschreibung)) + 1} {phase}", 2)
            doc.add_paragraph(beschreibung)
            
            # Leistungen für diese Phase finden
            leistungen = project_data.get("leistungen", [])
            phase_leistungen = [l for l in leistungen if phase in l.get("sia_phase", "")]
            
            if phase_leistungen:
                doc.add_paragraph("Enthaltene Leistungen:")
                for leistung in phase_leistungen:
                    doc.add_paragraph(f"  • {leistung.get('beschreibung', '-')}", style='List Bullet')
        
        # Projektumfang
        doc.add_heading("3. Projektumfang", 1)
        
        raeume = project_data.get("raeume", [])
        anlagen = project_data.get("anlagen", [])
        geraete = project_data.get("geraete", [])
        
        p = doc.add_paragraph()
        p.add_run(f"Anzahl Räume: {len(raeume)}").bold = True
        
        p = doc.add_paragraph()
        p.add_run(f"Anzahl Anlagen: {len(anlagen)}").bold = True
        
        p = doc.add_paragraph()
        p.add_run(f"Anzahl Geräte: {len(geraete)}").bold = True
        
        # Kalkulation (vereinfacht)
        doc.add_heading("4. Kalkulation", 1)
        doc.add_paragraph("Die Kalkulation basiert auf dem Projektumfang und den SIA-Phasen.")
        doc.add_paragraph("Detaillierte Kostenaufstellung auf Anfrage.")
        
        # Speichern
        filename = f"offerte_{project.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        
        return str(filepath)
