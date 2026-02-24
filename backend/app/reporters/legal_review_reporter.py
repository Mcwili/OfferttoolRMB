"""
Rechtliche Prüfung Reporter
Generiert Word-Dokument mit kritischen Paragraphen aus AI-Analyse
"""

from typing import Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import os
import re


class LegalReviewReporter:
    """Reporter für rechtliche Prüfung"""
    
    def __init__(self):
        # Pfad zur Vorlage: backend/Vorlagen (Deployment) oder Projekt-Root/Vorlagen (lokal)
        backend_dir = Path(__file__).resolve().parent.parent.parent  # Von reporters -> app -> backend
        project_root = backend_dir.parent  # Projekt-Root (ein Verzeichnis höher)
        
        template_paths = [
            backend_dir / "data" / "Vorlagen" / "RMB A4 hoch.docx",  # Deployment: /app/data/Vorlagen
            backend_dir / "Vorlagen" / "RMB A4 hoch.docx",
            project_root / "Vorlagen" / "RMB A4 hoch.docx",
            Path("Vorlagen") / "RMB A4 hoch.docx",
            Path("../Vorlagen") / "RMB A4 hoch.docx",
            Path("../../Vorlagen") / "RMB A4 hoch.docx",
        ]
        
        # Finde den ersten existierenden Pfad
        self.template_path = None
        for path in template_paths:
            if path.exists():
                self.template_path = path
                break
        
        if self.template_path is None:
            self.template_path = backend_dir / "data" / "Vorlagen" / "RMB A4 hoch.docx"
    
    def _format_general_assessment(self, doc: Document, text: str):
        """
        Formatiert die allgemeine Einschätzung strukturiert mit Aufzählungszeichen
        
        Args:
            doc: Word-Dokument
            text: Text der allgemeinen Einschätzung
        """
        if not text or not text.strip():
            return
        
        import re
        
        # Entferne überflüssige Leerzeichen und normalisiere
        text = text.strip()
        
        # Wenn der Text sehr lang ist (>400 Zeichen), kürze ihn
        max_length = 400
        if len(text) > max_length:
            # Versuche, bei Satzgrenzen zu kürzen
            shortened = text[:max_length]
            # Suche nach dem letzten Punkt, Ausrufezeichen oder Fragezeichen
            last_sentence_end = max(
                shortened.rfind('.'),
                shortened.rfind('!'),
                shortened.rfind('?')
            )
            if last_sentence_end > max_length * 0.7:  # Mindestens 70% des Textes behalten
                text = shortened[:last_sentence_end + 1]
            else:
                text = shortened + "..."
        
        # Teile den Text in Sätze auf
        sentences = re.split(r'[.!?]+\s+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        # Wenn der Text bereits strukturiert ist (mit Aufzählungszeichen), behalte die Struktur
        lines = text.split('\n')
        has_bullets = any(
            line.strip().startswith(('-', '•', '*', '·')) or
            (len(line.strip()) > 2 and line.strip()[0].isdigit() and line.strip()[1] in ['.', ')', ':'])
            for line in lines if line.strip()
        )
        
        if has_bullets:
            # Verwende die vorhandene Struktur
            paragraphs = []
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Entferne vorhandene Aufzählungszeichen
                cleaned = re.sub(r'^[-•*·]\s*', '', line)
                cleaned = re.sub(r'^\d+[.)]\s*', '', cleaned)
                
                if cleaned:
                    paragraphs.append(cleaned)
            
            # Begrenze auf maximal 5 Punkte für bessere Lesbarkeit
            paragraphs = paragraphs[:5]
        else:
            # Erstelle strukturierte Liste aus Sätzen
            # Gruppiere verwandte Sätze zusammen
            paragraphs = []
            current_group = []
            
            for sentence in sentences:
                # Wenn der Satz sehr kurz ist (<30 Zeichen), füge ihn zur aktuellen Gruppe hinzu
                if len(sentence) < 30 and current_group:
                    current_group.append(sentence)
                else:
                    # Speichere vorherige Gruppe
                    if current_group:
                        paragraphs.append(' '.join(current_group))
                    # Starte neue Gruppe
                    current_group = [sentence]
            
            # Füge letzte Gruppe hinzu
            if current_group:
                paragraphs.append(' '.join(current_group))
            
            # Begrenze auf maximal 4-5 Punkte für bessere Lesbarkeit
            if len(paragraphs) > 5:
                # Kombiniere die letzten Punkte
                combined = ' '.join(paragraphs[4:])
                paragraphs = paragraphs[:4] + [combined]
        
        # Formatiere als strukturierte Liste mit Aufzählungszeichen
        for para_text in paragraphs:
            if not para_text:
                continue
            
            para_text = para_text.strip()
            
            # Entferne doppelte Leerzeichen
            para_text = re.sub(r'\s+', ' ', para_text)
            
            # Kürze sehr lange Punkte (max. 200 Zeichen pro Punkt)
            if len(para_text) > 200:
                para_text = para_text[:197] + "..."
            
            # Erstelle Absatz mit Aufzählungszeichen
            # Verwende immer manuelle Formatierung, da der Style möglicherweise nicht existiert
            p = doc.add_paragraph()
            # Füge Bullet Point hinzu
            bullet_run = p.add_run('• ')
            bullet_run.font.size = Pt(11)
            text_run = p.add_run(para_text)
            text_run.font.size = Pt(11)
            # Einrückung für Aufzählung
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
    
    async def generate(self, project_name: str, analysis_result: Dict[str, Any]) -> bytes:
        """
        Generiert Word-Dokument mit rechtlicher Prüfung
        
        Args:
            project_name: Name des Projekts
            analysis_result: Ergebnis der AI-Analyse mit allgemeine_einschaetzung und kritische_punkte
        
        Returns:
            Bytes des Word-Dokuments
        """
        import logging
        logger = logging.getLogger(__name__)
        
        # Logging: Prüfe, welche Daten empfangen werden
        logger.info(f"Reporter empfängt Daten: {len(analysis_result.get('kritische_punkte', []))} kritische Punkte")
        logger.debug(f"Allgemeine Einschätzung Länge: {len(analysis_result.get('allgemeine_einschaetzung', ''))} Zeichen")
        if analysis_result.get('kritische_punkte'):
            logger.debug(f"Erster kritischer Punkt: {analysis_result['kritische_punkte'][0].get('titel', 'Kein Titel')}")
            logger.debug(f"Letzter kritischer Punkt: {analysis_result['kritische_punkte'][-1].get('titel', 'Kein Titel')}")
        # Versuche Vorlage zu laden, sonst neues Dokument
        if self.template_path and self.template_path.exists():
            try:
                # Lade Vorlage
                doc = Document(str(self.template_path))
                
                # Entferne ALLE Absätze aus der Vorlage (behält Formatierung)
                # Dies verhindert leere Seiten
                paragraphs_to_remove = list(doc.paragraphs)
                for para in paragraphs_to_remove:
                    p_element = para._element
                    p_element.getparent().remove(p_element)
                
                # Entferne alle Seitenumbrüche aus dem Dokument
                from docx.oxml.ns import qn
                try:
                    # Suche nach allen Seitenumbruch-Elementen
                    body = doc.element.body
                    elements_to_remove = []
                    
                    for element in body:
                        # Entferne direkte Seitenumbrüche
                        if 'br' in element.tag:
                            if element.get(qn('w:type')) == 'page':
                                elements_to_remove.append(element)
                        # Entferne Seitenumbrüche in Absätzen
                        elif 'p' in element.tag:
                            # Suche nach Seitenumbrüchen innerhalb des Absatzes
                            br_elements = element.findall('.//w:br', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                            for br in br_elements:
                                if br.get(qn('w:type')) == 'page':
                                    br.getparent().remove(br)
                    
                    # Entferne gefundene Elemente
                    for element in elements_to_remove:
                        body.remove(element)
                except Exception as e:
                    print(f"Hinweis: Konnte nicht alle Seitenumbrüche entfernen: {e}")
            except Exception as e:
                # Falls Fehler beim Laden der Vorlage, verwende leeres Dokument
                print(f"Warnung: Vorlage konnte nicht geladen werden: {e}")
                doc = Document()
        else:
            doc = Document()
        
        # Titel (linksbündig)
        title = doc.add_heading(f"Rechtliche Prüfung: {project_name}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Datum
        p = doc.add_paragraph()
        p.add_run("Erstellt am: ").bold = True
        p.add_run(datetime.now().strftime("%d.%m.%Y %H:%M"))
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Allgemeine Einschätzung
        doc.add_heading("Allgemeine Einschätzung", 1)
        allgemeine = analysis_result.get("allgemeine_einschaetzung", "")
        logger.debug(f"Allgemeine Einschätzung für Word: {len(allgemeine)} Zeichen")
        if allgemeine:
            # Formatiere den Text strukturiert mit Aufzählungszeichen
            self._format_general_assessment(doc, allgemeine)
        else:
            doc.add_paragraph("Keine allgemeine Einschätzung verfügbar.")
        
        # Kritische Punkte
        kritische_punkte = analysis_result.get("kritische_punkte", [])
        logger.info(f"Kritische Punkte für Word: {len(kritische_punkte)} Punkte")
        if kritische_punkte:
            doc.add_heading("Kritische Punkte", 1)
            
            for idx, punkt in enumerate(kritische_punkte):
                # Logging für ersten und letzten Punkt
                if idx == 0:
                    logger.debug(f"Erster Punkt im Word: [{punkt.get('nummer', '?')}] {punkt.get('titel', 'Unbekannter Titel')}")
                if idx == len(kritische_punkte) - 1:
                    logger.debug(f"Letzter Punkt im Word: [{punkt.get('nummer', '?')}] {punkt.get('titel', 'Unbekannter Titel')}")
                
                # Nummer und Titel
                heading_text = f"[{punkt.get('nummer', '?')}] {punkt.get('titel', 'Unbekannter Titel')}"
                doc.add_heading(heading_text, 2)
                
                # Zitat mit Quellenangabe direkt beim Zitat
                zitat = punkt.get("zitat", "")
                quelle_datei = punkt.get("quelle_datei")
                quelle_paragraph = punkt.get("quelle_paragraph")
                
                if zitat:
                    p = doc.add_paragraph()
                    p.add_run("Zitat: ").bold = True
                    zitat_run = p.add_run(zitat)
                    zitat_run.italic = True
                    
                    # Quelle direkt nach dem Zitat hinzufügen
                    if quelle_datei or quelle_paragraph:
                        quelle_text_parts = []
                        if quelle_datei:
                            quelle_text_parts.append(quelle_datei)
                        if quelle_paragraph is not None:
                            quelle_text_parts.append(f"Absatz {quelle_paragraph}")
                        
                        if quelle_text_parts:
                            source_text = " (Quelle: " + ", ".join(quelle_text_parts) + ")"
                            source_run = p.add_run(source_text)
                            source_run.italic = False
                            source_run.font.color.rgb = RGBColor(0, 122, 255)  # Blau für Quelle
                    
                    # Einrückung für Zitat
                    p.paragraph_format.left_indent = Inches(0.5)
                
                # Beurteilung
                beurteilung = punkt.get("beurteilung", "")
                if beurteilung:
                    p = doc.add_paragraph()
                    p.add_run("Beurteilung: ").bold = True
                    p.add_run(beurteilung)
                
                # Risiko Rating mit Farbe
                rating = punkt.get("risiko_rating", "").lower()
                p = doc.add_paragraph()
                p.add_run("Risiko Rating: ").bold = True
                rating_run = p.add_run(rating.upper())
                
                # Farben basierend auf Rating
                if rating == "rot":
                    rating_run.font.color.rgb = RGBColor(255, 0, 0)  # Rot
                elif rating == "orange":
                    rating_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
                elif rating == "grün":
                    rating_run.font.color.rgb = RGBColor(0, 128, 0)  # Grün
                
                # Empfehlung
                empfehlung = punkt.get("empfehlung", "")
                if empfehlung:
                    p = doc.add_paragraph()
                    p.add_run("Empfehlung: ").bold = True
                    p.add_run(empfehlung)
                
                # Leerzeile zwischen Punkten
                doc.add_paragraph()
        else:
            doc.add_paragraph("Keine kritischen Punkte gefunden.")
        
        # Dokument als Bytes zurückgeben
        from io import BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()
