"""
Question List Reporter
Generiert Word-Dokument mit Frageliste aus AI-Analyse
"""

from typing import Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import os


class QuestionListReporter:
    """Reporter für Frageliste"""
    
    def __init__(self):
        # Pfad zur Vorlage: backend/Vorlagen (Deployment) oder Projekt-Root/Vorlagen (lokal)
        backend_dir = Path(__file__).resolve().parent.parent.parent  # Von reporters -> app -> backend
        project_root = backend_dir.parent  # Projekt-Root (ein Verzeichnis höher)
        
        template_paths = [
            backend_dir / "data" / "Vorlagen" / "Frageliste Vorlage.docx",  # Deployment: /app/data/Vorlagen
            backend_dir / "data" / "Vorlagen" / "RMB A4 hoch.docx",
            backend_dir / "Vorlagen" / "Frageliste Vorlage.docx",
            backend_dir / "Vorlagen" / "RMB A4 hoch.docx",
            project_root / "Vorlagen" / "Frageliste Vorlage.docx",
            project_root / "Vorlagen" / "RMB A4 hoch.docx",
            Path("Vorlagen") / "Frageliste Vorlage.docx",
            Path("Vorlagen") / "RMB A4 hoch.docx",
            Path("../Vorlagen") / "Frageliste Vorlage.docx",
            Path("../Vorlagen") / "RMB A4 hoch.docx",
        ]
        
        # Finde den ersten existierenden Pfad
        self.template_path = None
        for path in template_paths:
            if path.exists():
                self.template_path = path
                break
        
        if self.template_path is None:
            self.template_path = backend_dir / "data" / "Vorlagen" / "RMB A4 hoch.docx"
    
    async def generate(self, project_name: str, analysis_result: Dict[str, Any]) -> bytes:
        """
        Generiert Word-Dokument mit Frageliste
        
        Args:
            project_name: Name des Projekts
            analysis_result: Ergebnis der AI-Analyse mit zusammenfassung und fragen
        
        Returns:
            Bytes des Word-Dokuments
        """
        # Versuche Vorlage zu laden, sonst neues Dokument
        if self.template_path and self.template_path.exists():
            try:
                # Lade Vorlage
                doc = Document(str(self.template_path))
                
                # Entferne ALLE Absätze aus der Vorlage (behält Formatierung)
                # Dies verhindert leere Seiten
                paragraphs_to_remove = list(doc.paragraphs)
                for para in paragraphs_to_remove:
                    p_element = para._element
                    p_element.getparent().remove(p_element)
                
                # Entferne alle Seitenumbrüche aus dem Dokument
                from docx.oxml.ns import qn
                try:
                    # Suche nach allen Seitenumbruch-Elementen
                    body = doc.element.body
                    elements_to_remove = []
                    
                    for element in body:
                        # Entferne direkte Seitenumbrüche
                        if 'br' in element.tag:
                            if element.get(qn('w:type')) == 'page':
                                elements_to_remove.append(element)
                        # Entferne Seitenumbrüche in Absätzen
                        elif 'p' in element.tag:
                            # Suche nach Seitenumbrüchen innerhalb des Absatzes
                            br_elements = element.findall('.//w:br', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                            for br in br_elements:
                                if br.get(qn('w:type')) == 'page':
                                    br.getparent().remove(br)
                    
                    # Entferne gefundene Elemente
                    for element in elements_to_remove:
                        body.remove(element)
                except Exception as e:
                    print(f"Hinweis: Konnte nicht alle Seitenumbrüche entfernen: {e}")
            except Exception as e:
                # Falls Fehler beim Laden der Vorlage, verwende leeres Dokument
                print(f"Warnung: Vorlage konnte nicht geladen werden: {e}")
                doc = Document()
        else:
            doc = Document()
        
        # Titel (linksbündig)
        title = doc.add_heading(f"Frageliste: {project_name}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Datum
        p = doc.add_paragraph()
        p.add_run("Erstellt am: ").bold = True
        p.add_run(datetime.now().strftime("%d.%m.%Y %H:%M"))
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Zusammenfassung
        doc.add_heading("Zusammenfassung", 1)
        zusammenfassung = analysis_result.get("zusammenfassung", "")
        if zusammenfassung:
            doc.add_paragraph(zusammenfassung)
        else:
            doc.add_paragraph("Keine Zusammenfassung verfügbar.")
        
        # Fragen nach Kategorien gruppiert
        fragen = analysis_result.get("fragen", [])
        if fragen:
            doc.add_heading("Fragen", 1)
            
            # Gruppiere Fragen nach Kategorien
            fragen_by_kategorie = {}
            for frage in fragen:
                kategorie = frage.get("kategorie", "Sonstiges")
                if kategorie not in fragen_by_kategorie:
                    fragen_by_kategorie[kategorie] = []
                fragen_by_kategorie[kategorie].append(frage)
            
            # Iteriere durch Kategorien
            for kategorie, kategorie_fragen in fragen_by_kategorie.items():
                # Kategorie als Heading
                doc.add_heading(kategorie, 2)
                
                # Fragen dieser Kategorie
                for frage in kategorie_fragen:
                    # Nummer und Frage
                    nummer = frage.get("nummer", "?")
                    frage_text = frage.get("frage", "")
                    
                    p = doc.add_paragraph()
                    p.add_run(f"Frage {nummer}: ").bold = True
                    p.add_run(frage_text)
                    p.paragraph_format.left_indent = Inches(0.3)
                    
                    # Begründung
                    begruendung = frage.get("begruendung", "")
                    if begruendung:
                        p = doc.add_paragraph()
                        p.add_run("Begründung: ").bold = True
                        p.add_run(begruendung)
                        p.paragraph_format.left_indent = Inches(0.5)
                    
                    # Priorität mit Farbe
                    prioritaet = frage.get("prioritaet", "").lower()
                    p = doc.add_paragraph()
                    p.add_run("Priorität: ").bold = True
                    prioritaet_run = p.add_run(prioritaet.upper())
                    
                    # Farben basierend auf Priorität
                    if prioritaet == "hoch":
                        prioritaet_run.font.color.rgb = RGBColor(255, 0, 0)  # Rot
                    elif prioritaet == "mittel":
                        prioritaet_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
                    elif prioritaet == "niedrig":
                        prioritaet_run.font.color.rgb = RGBColor(0, 128, 0)  # Grün
                    
                    p.paragraph_format.left_indent = Inches(0.5)
                    
                    # Leerzeile zwischen Fragen
                    doc.add_paragraph()
        else:
            doc.add_paragraph("Keine Fragen gefunden.")
        
        # Dokument als Bytes zurückgeben
        from io import BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()
